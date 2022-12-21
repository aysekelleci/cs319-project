import docx
import pytz
import os
from datetime import datetime

from django.shortcuts import render
from django.views import View
from django.shortcuts import get_object_or_404, render
from django.shortcuts import redirect
from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.files import File

from courses.forms import CourseForm, DocumentForm, CoordinatorDocumentForm, BilkentCourseForm

from .models import Course, Document, MergedCourse, MUST_COURSE, ELECTIVE_COURSE, PREAPPROVAL_FORM, STATIC_DOCUMENTS_FOLDER
from accounts.models import UserCourse, Student, ErasmusUser, Coordinator, ToDo, BoardMember
from accounts.models import INITIAL, PLACED, NO_PLACEMENT, CHOOSING_COURSES, WAIT_COURSE_APPROVAL, WAIT_FINAL_LIST_APPROVAL,\
    FINAL_LIST_APPROVED, WAIT_PRE_APPROVAL_FORM, WAIT_MOBILITY, IN_MOBILITY, FINISHED_MOBILITY
from communication.models import Notification

import os
from django.conf import settings
from django.http import HttpResponse, Http404
from abc import ABC, abstractmethod

def sendNotification(header, user, link=""):
    new_notif = Notification(header=header, link=link, user=user) # user: erasmus_user
    new_notif.save()

def sendTodo(header, user, body='', link=''):
    new_todo = ToDo(header=header, body=body, link=link, user=user)
    new_todo.save()


def download(request, path):
    # get the download path
    download_path = os.path.join(settings.MEDIA_ROOT, path)
    if os.path.exists(download_path):

        with open(download_path, 'rb') as fh:

            response = HttpResponse(fh.read(), content_type="application/document")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(download_path)
            return response
    raise Http404

#######################################################################################################################
# Course views


class CourseView(LoginRequiredMixin,View):
    def get(self, request):
        student = None
        user = request.user
        approved_unmerged_courses = None
        approved_merged_course_dict = None
        user_unmerged_courses = None
        user_merged_course_dict = None
        rejected_courses = None
        total_ects_credit = None
        erasmus_user = ErasmusUser.objects.get(user=user)
        if Coordinator.objects.filter(user=erasmus_user).first():
            user_type = "Coordinator"

        elif Student.objects.filter(user=erasmus_user).first():
            user_type = "Student"
            student = Student.objects.filter(user=erasmus_user).first()
            # get preapproved courses for the university the student will be attending
            _approved_courses = Course.objects.filter(university=student.university, approved=True)
            _approved_merged_courses = {course.merged_course for course in _approved_courses if course.is_merged is True}
            approved_merged_course_dict = {}

            for merged_course in _approved_merged_courses:
                one_merged_course_contents = []  # each list contains the courses composing one merged course
                for course in _approved_courses:
                    if (course.is_merged and (course.merged_course.pk is merged_course.pk)):
                        one_merged_course_contents.append(course)
                approved_merged_course_dict[merged_course] = one_merged_course_contents

            approved_unmerged_courses = [course for course in _approved_courses if course.is_merged is False]

            _user_courses = UserCourse.objects.filter(user=student)
            total_ects_credit = sum([course.course.course_credit for course in _user_courses])
            _merged_courses = {course.course.merged_course for course in _user_courses if course.course.is_merged is True}
            user_merged_course_dict = getMergedCoursesDict(_user_courses, _merged_courses)
            user_unmerged_courses = [user_course for user_course in _user_courses if user_course.course.is_merged is False]

            rejected_courses = Course.objects.filter(university=student.university, is_rejected=True)
        else:
            user_type = "Board Member"

        context = {'user': user, 'approved_unmerged_courses': approved_unmerged_courses,
                   "approved_merged_course_dict": approved_merged_course_dict, "user_type": user_type, 'student': student,
                   'user_unmerged_courses': user_unmerged_courses, 'user_merged_course_dict': user_merged_course_dict,
                   'rejected_courses': rejected_courses, 'total_ects_credit': total_ects_credit,'MUST_COURSE': MUST_COURSE,
                   'ELECTIVE_COURSE': ELECTIVE_COURSE}

        return render(request, 'courses/courses.html', context)
def getMergedCoursesDict(user_courses, merged_courses):
    merged_course_dict = {}
    for merged_course in merged_courses:
        one_merged_course_contents = []  # each list contains the courses composing one merged course
        for user_course in user_courses:
            if (user_course.course.is_merged and (user_course.course.merged_course.pk is merged_course.pk) ):
                one_merged_course_contents.append(user_course)
        merged_course_dict[merged_course] = one_merged_course_contents
    return merged_course_dict


class AddCourseView(LoginRequiredMixin, View):
    def get(self, request, course_id):
        course = get_object_or_404(Course, pk=course_id)  # Check whether given course object exists
        user = request.user
        erasmus_user = ErasmusUser.objects.get(user=user)
        student = Student.objects.get(user=erasmus_user)  # get which student wants to add approved course
        try:
            course_item = UserCourse.objects.get(user=student,
                                                 course=course)  # control whether course item already added
        except:
            course_item = None
        if course_item is not None:
            messages.error(request, "This course already added to course list.")
            return redirect("/courses")

        else:
            course = UserCourse.objects.create(user=student, course=course)
            messages.success(request, "This course was added to your course list.")
            return redirect("/courses")


class DeleteCourseView(LoginRequiredMixin,View):
    def get(self, request, course_id):
        user = request.user
        erasmus_user = ErasmusUser.objects.get(user=user)
        student = Student.objects.get(user=erasmus_user)  # get which student wants to add approved course
        course_to_delete = get_object_or_404(Course, pk=course_id)

        user_course_to_delete = UserCourse.objects.filter(user=student, course=course_to_delete).first()

        if user_course_to_delete is not None:
            if course_to_delete.is_merged:
                user_courses = UserCourse.objects.filter(user=student)
                for user_course in user_courses:
                    course = user_course.course
                    if course.is_merged and course.merged_course.pk is course_to_delete.merged_course.pk:
                        # change associated merged courses to not be merged
                        course.is_merged = False
                        course.save()
                        # delete all merged user courses
                        user_course.delete()
                # delete the parent merged course object
                course_to_delete.merged_course.delete()
                messages.success(request, "Merged courses removed from the list.")
            else:
                user_course_to_delete.delete()
                messages.success(request, "Course removed from the list.")
        return redirect("/courses")


class AddUnapprovedCourse(LoginRequiredMixin, View):

    def get(self, request):
        username = request.user.username

        new_course = None

        course_form = CourseForm()

        return render(request,
                      'courses/add_unapproved_course.html',
                      {'new_course': new_course,
                       'course_form': course_form,
                       'username': username})

    def post(self, request):
        user = request.user
        username = user.username
        erasmus_user = ErasmusUser.objects.get(user=user)
        student = Student.objects.filter(user=erasmus_user).first()
        university = student.university

        course_form = CourseForm(data=request.POST)
        if course_form.is_valid():
            # Create course object but don't save to database yet
            new_course = course_form.save(commit=False)

            new_course.university = university

            # Save the course to the database
            new_course.save()

            # Create a course specific for the user as well
            new_user_course = UserCourse(course=new_course, user=student)
            new_user_course.save()

        else:
            messages.info(request, "Course Form is not valid")
            return redirect("add_unapproved_course")
        return redirect("/courses")


class AddBilkentCourse(LoginRequiredMixin, View):
    def get(self, request):
        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        bilkent_course_form = BilkentCourseForm()
        context = {'erasmus_user': erasmus_user, 'bilkent_course_form': bilkent_course_form}
        return render(request, 'courses/add_bilkent_course.html', context)

    def post(self, request):
        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        new_course = None

        bilkent_course_form = BilkentCourseForm(data=request.POST)
        if bilkent_course_form.is_valid():
            new_course = bilkent_course_form.save(commit=False)
            if new_course.course_type == MUST_COURSE:
                if new_course.course_name == '' or  new_course.course_code == '':
                    context = {'erasmus_user': erasmus_user, 'bilkent_course_form': bilkent_course_form,
                               'new_course': new_course}
                    messages.error(request, 'You must write course name and course code for must courses')
                    return render(request, 'courses/add_bilkent_course.html', context)
            # Create course object but don't save to database yet
            if new_course.course_type == ELECTIVE_COURSE:
                if new_course.elective_group_name == '':
                    context = {'erasmus_user': erasmus_user, 'bilkent_course_form': bilkent_course_form,
                               'new_course': new_course}
                    messages.error(request, 'You must select elective group name for elective courses')
                    return render(request, 'courses/add_bilkent_course.html', context)

            # Save the course to the database
            new_course.save()
        else:
            messages.error(request,'Bilkent course form is not valid')
            context = {'erasmus_user': erasmus_user, 'bilkent_course_form': bilkent_course_form, 'new_course': new_course}
            return render(request, 'courses/add_bilkent_course.html', context)

        messages.success(request, 'Bilkent course is added to course list successfully')
        return redirect('/courses')

class GetWaitingCoursesView(LoginRequiredMixin, View):
    def get(self, request):
        unapproved_unmerged_courses = UserCourse.objects.filter(course__approved__exact=False,
                                                                course__is_rejected__exact=False,
                                                                submitted__exact=True,
                                                                course__is_merged__exact=False)

        unapproved_merged_courses = UserCourse.objects.all().filter(course__merged_course__approved__exact=False,
                                                             course__merged_course__is_rejected__exact=False,
                                                             course__merged_course__submitted__exact=True,
                                                             course__is_merged__exact=True).distinct("course__merged_course")

        unapproved_final_list_students = Student.objects.all().filter(status__exact=WAIT_FINAL_LIST_APPROVAL)

        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        coordinator = Coordinator.objects.filter(user=erasmus_user).first()

        if coordinator is None:
            messages.error(request, "Only the coordinator can see waiting approvals.")
            redirect("accounts/profile")
        else:
            context = {'coordinator': coordinator, 'unapproved_unmerged_courses': unapproved_unmerged_courses,
                       'unapproved_merged_courses': unapproved_merged_courses,
                       'unapproved_final_list_students': unapproved_final_list_students}
            return render(request, 'courses/waiting_courses.html', context)

class SubmitCourseView(LoginRequiredMixin, View):
    def get(self, request, course_id):
        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        student = Student.objects.filter(user=erasmus_user).first()

        user_course = UserCourse.objects.filter(course__pk__exact=course_id).first()

        # update user course's submitted status
        if user_course.course.is_merged:
            user_course.course.merged_course.submitted = True
            user_course.course.merged_course.save()
        else:
            user_course.submitted = True
            user_course.save()

        # update student's status
        student.status = WAIT_COURSE_APPROVAL
        student.save()

        # send notification and todo to the coordinator fixme does the link work?
        # send notification and todo to the coordinator
        sendNotification(header=f"{erasmus_user.name} submitted a course for your approval.",
                         user=student.coordinator.user)
        sendTodo(header=f"Approve {erasmus_user.name}'s submitted course.", user=student.coordinator.user)


        messages.success(request, "Course submitted for approval.")
        return redirect('/courses')

class SubmitCourseListView(LoginRequiredMixin, View):
    def get(self, request):
        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        student = Student.objects.filter(user=erasmus_user).first()

        user_courses = UserCourse.objects.all().filter(user=student)

        # check if the user courses are all approved (merged or unmerged)
        for user_course in user_courses:
            if (user_course.course.is_merged and not user_course.course.merged_course.approved) \
                    or (not user_course.course.is_merged and not user_course.course.approved):
                messages.error(request, "You can only submit the final list only if all the courses are approved.")
                return redirect('/courses')

        # update student's status
        student.final_list_submitted = True
        student.status = WAIT_FINAL_LIST_APPROVAL
        student.save()

        sendNotification(header=f"{erasmus_user.name} submitted their final course list for your approval.",
                         link=f"/profile/{student.pk}",
                         user=student.coordinator.user)
        sendTodo(header=f"Approve {erasmus_user.name}'s final course list.", user=student.coordinator.user)

        messages.success(request, "Final list submitted.")
        return redirect('/courses')

class ApproveCoursesView(LoginRequiredMixin, View):
    def get(self, request, course_id, student_id):
        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        coordinator = Coordinator.objects.filter(user=erasmus_user).first()
        student = Student.objects.filter(pk=student_id).first()
        if coordinator is None:
            messages.error(request, "Only coordinators can evaluate courses.")
            return redirect("accounts/profile")
        else:
            course = Course.objects.filter(pk=course_id).first()
            if course is not None:
                if course.is_merged:
                    course.merged_course.approved = True
                    course.merged_course.save()
                else:
                    course.approved = True
                course.save()
                messages.success(request, "Course approved.")
                sendNotification(header=f"{coordinator.user.name} approved your course.",
                                 user=student.user)
            else:
                messages.error(request, "This is not a course.")

            return redirect(request.META.get('HTTP_REFERER', 'redirect_if_referer_not_found'))


class RejectCourseView(LoginRequiredMixin, View):
    def get(self, request, course_id, student_id):
        unapproved_courses = Course.objects.filter(approved=False)
        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        coordinator = Coordinator.objects.filter(user=erasmus_user).first()
        student = Student.objects.filter(pk=student_id).first()
        if coordinator is None:
            messages.error(request, "Only coordinators can evaluate courses.")
            return redirect('accounts/profile')
        else:
            course = Course.objects.filter(pk=course_id).first()
            if course is not None:
                if course.is_merged:
                    course.merged_course.is_rejected = True
                    course.merged_course.save()
                else:
                    course.is_rejected = True
                course.save()
                messages.success(request, "Course rejected.")
                sendNotification(header=f"{coordinator.user.name} rejected your course.",
                                 user=student.user)
            else:
                messages.error(request, "This is not a course.")

            return redirect(request.META.get('HTTP_REFERER', 'redirect_if_referer_not_found'))

class ApproveFinalListView(LoginRequiredMixin,View):
    def get(self, request, student_id):

        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        coordinator = Coordinator.objects.filter(user=erasmus_user).first()
        if coordinator is None:
            messages.error(request, "Only coordinators can evaluate courses.")
            return redirect("accounts/profile")

        # get student
        student = Student.objects.filter(pk=student_id).first()

        # check that the student has at least one course
        user_courses = UserCourse.objects.all().filter(user=student)
        if user_courses is None:
            messages.error(request, "User should have at least one course.")
            return redirect(request.META.get('HTTP_REFERER', 'redirect_if_referer_not_found'))

        # check if the user courses are all approved (merged or unmerged)
        for user_course in user_courses:
            if (user_course.course.is_merged and not user_course.course.merged_course.approved) or (
            not user_course.course.approved):
                messages.error(request, "All courses should be approved first to evaluate the final list.")
                return redirect(request.META.get('HTTP_REFERER', 'redirect_if_referer_not_found'))

        # update student's status
        student.final_list_approved = True
        student.status = FINAL_LIST_APPROVED
        student.save()
        sendNotification(header=f"{coordinator.user.name} approved your course list.",
                         user=student.user)

        return redirect(request.META.get('HTTP_REFERER', 'redirect_if_referer_not_found'))


class RejectFinalListView(LoginRequiredMixin, View):
    def get(self, request, student_id):

        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        coordinator = Coordinator.objects.filter(user=erasmus_user).first()
        if coordinator is None:
            messages.error(request, "Only coordinators can evaluate courses.")
            return redirect("accounts/profile")

        # get student
        student = Student.objects.filter(pk=student_id).first()

        # check that the student has at least one course
        user_courses = UserCourse.objects.all().filter(user=student)
        if user_courses is None:
            messages.error(request, "User should have at least one course.")
            return redirect(request.META.get('HTTP_REFERER', 'redirect_if_referer_not_found'))

        # check if the user courses are all approved (merged or unmerged)
        for user_course in user_courses:
            if (user_course.course.is_merged and not user_course.course.merged_course.approved) or (
                    not user_course.course.approved):
                messages.error(request, "All courses should be approved first to evaluate the final list.")
                return redirect(request.META.get('HTTP_REFERER', 'redirect_if_referer_not_found'))

        # change student's status
        student.final_list_submitted = False
        student.status = CHOOSING_COURSES
        student.save()

        sendNotification(header=f"{coordinator.user.name} rejected your course list.",
                         user=student.user)

        return redirect(request.META.get('HTTP_REFERER', 'redirect_if_referer_not_found'))
class MergeCourseView(LoginRequiredMixin, View):
    def get(self, request, course_id1, course_id2, course_id3, course_id4, course_id5,
            course_id6, course_id7, course_id8, course_id9, course_id10):

        if Course.objects.filter(pk=course_id1).exists() and Course.objects.filter(pk=course_id2).exists():
            # get the bilkent equivalent course
            bilkent_equivalent = Course.objects.filter(pk=course_id1).first().bilkent_equivalent

            course_ids = [course_id1, course_id2, course_id3, course_id4, course_id5, course_id6, course_id7,
                          course_id8, course_id9, course_id10]
            # check if Bilkent equivalent courses of all courses are the same
            for course_id in course_ids:
                course = Course.objects.filter(pk=course_id).first()
                if course is not None and course.bilkent_equivalent.pk is not bilkent_equivalent.pk:
                    messages.error(request, "Bilkent equivalent courses of all courses to be merged must be the same.")
                    return redirect("/courses")

            # create a "parent" merged course instance
            merged_course = MergedCourse(bilkent_equivalent=bilkent_equivalent)
            merged_course.save()

            # merge up to 10 courses by designating their parent course as the created MergedCourse instance
            for course_id in course_ids:
                course = Course.objects.filter(pk=course_id).first()
                if course is not None:
                    course.is_merged = True
                    course.merged_course = merged_course
                    course.save()

            messages.success(request, "Successfully merged courses.")
            return redirect("/courses")
        else:
            messages.error(request, "You need to choose at least two courses to merge.")
            return redirect("/courses")

#######################################################################################################################
# Document  views

class DocumentView(LoginRequiredMixin, View):
    def get(self, request):
        documents = Document.objects.all().order_by('-date')
        user = request.user

        context = {'documents': documents}
        return render(request, 'courses/documents.html', context)


class UploadDocumentView(LoginRequiredMixin, View):
    def get(self, request, from_student_profile, viewed_student_id):
        user = request.user
        student = None
        coordinator = None
        document_form = None
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        if Student.objects.filter(user=erasmus_user).exists():
            # Student wants to add a document
            student = Student.objects.filter(user=erasmus_user).first()
            document_form = DocumentForm()
        elif Coordinator.objects.filter(user=erasmus_user).exists():
            coordinator = Coordinator.objects.filter(user=erasmus_user).first()
            if from_student_profile:
                # Coordinator wants to add a document from a student's profile page
                document_form = DocumentForm()
            else:
                # Coordinator wants to add a document from their own files page
                document_form = CoordinatorDocumentForm()
        else:
            redirect("/login")

        context = {'student': student, 'coordinator': coordinator, 'document_form': document_form}
        return render(request, 'courses/upload-documents.html', context)

    def post(self, request, from_student_profile, viewed_student_id):
        student = None
        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        # for the student:
        if Student.objects.filter(user=erasmus_user).exists():
            student = Student.objects.filter(user=erasmus_user).first()
            document_form = DocumentForm(request.POST, request.FILES)

            if document_form.is_valid():
                new_document = document_form.save(commit=False)
                new_document.user = student
                # Save the document to the database
                new_document.save()
                messages.success(request, "Document is added")
                return redirect("/documents")
            else:
                messages.error(request, "Document form is not valid")
                return redirect("/documents")
        # for the coordinator:
        elif Coordinator.objects.filter(user=erasmus_user).exists():
            if from_student_profile:
                document_form = DocumentForm(request.POST, request.FILES)
            else:
                document_form = CoordinatorDocumentForm(request.POST, request.FILES)

            if document_form.is_valid():
                new_document = document_form.save(commit=False)
                new_document.is_signed_coordinator = True # It is assumed that the coordinator would only add signed documents
                if from_student_profile:
                    new_document.user = Student.objects.filter(pk=viewed_student_id).first()
                # Save the document to the database
                new_document.save()
                messages.success(request, "Document is added")
                new_document.student.status = WAIT_MOBILITY
                return redirect("/documents")
            else:
                messages.error(request, "Document form is not valid")
                return redirect("/documents")
        else:
            redirect("/login")


class DownloadDocument(View):
    def get(self, request, document_id):
        document = get_object_or_404(Document, pk=document_id)


class DeleteDocumentView(LoginRequiredMixin, View):
    def get(self, request, document_id):
        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        if erasmus_user is not None:
            student = Student.objects.filter(user=erasmus_user).first()
            document = get_object_or_404(Document, pk=document_id, user=student)

            document.delete()

        return redirect("/documents")


class CreateDocumentView(LoginRequiredMixin, View, ABC):
    @abstractmethod
    def fill_necessary_information(self, student):
        pass
    def get(self, request):
        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        student = Student.objects.filter(user=erasmus_user).first()

        if student.university is None:
            messages.error(request, "You need to be placed in a university to generate documents.")
            return redirect("/courses")

        document_name, date, document_type = self.fill_necessary_information(student)

        with open(document_name, 'rb') as f:
            new_pre_approval = Document(document_name='Unsigned ' + document_type.lower() + ' of ' + erasmus_user.name,
                                        date=date, is_signed=False, user=student,
                                        document_type=document_type)
            new_pre_approval.document = File(f, name=os.path.basename(f.name))
            new_pre_approval.save()
            messages.success(request, "Document is generated")
            student.status = WAIT_PRE_APPROVAL_FORM
            # send notification and todo to the coordinator
            sendNotification(header=f"{student.user.name} created a {document_type}.",
                             user=student.user)
            sendTodo(header=f"Sign {student.user.name}'s {document_type}.", user=student.coordinator.user)
            student.save()

        return redirect("/courses")

class CreatePreApprovalView(CreateDocumentView):
    def fill_necessary_information(self, student):
        # read the template pre-approval form and its tables
        document = docx.api.Document('courses/static/pre_approval_form.docx')
        student_info_table = document.tables[0]
        host_uni_table = document.tables[1]
        courses_table = document.tables[2]
        coordinator_table = document.tables[3]

        student_info_table.cell(0, 1).text = student.user.name.split()[:-1]  # name
        student_info_table.cell(0, 3).text = str(student.user.bilkent_id)  # bilkent id number
        student_info_table.cell(1, 1).text = student.user.name.split()[-1]  # surname
        student_info_table.cell(1, 3).text = student.user.department  # department

        host_uni_table.cell(0, 1).text = student.university.university_name  # host university name
        host_uni_table.cell(0, 3).text = student.academic_year  # academic year
        host_uni_table.cell(1, 3).text = student.semester  # semester

        courses = UserCourse.objects.filter(user=student)

        unmerged_courses = UserCourse.objects.filter(user=student, course__is_merged__exact = False)
        _user_courses = UserCourse.objects.filter(user=student)
        _merged_courses = {course.course.merged_course for course in _user_courses if course.course.is_merged is True}
        user_merged_course_dict = getMergedCoursesDict(_user_courses, _merged_courses)


        for i, user_course in enumerate(unmerged_courses):
            course = user_course.course
            bilkent_equivalent = course.bilkent_equivalent
            courses_table.cell(2 + i, 1).text = course.code  # course code
            courses_table.cell(2 + i, 2).text = course.course_name  # course name
            courses_table.cell(2 + i, 3).text = str(course.course_credit)  # course credit
            if bilkent_equivalent is not None:
                if bilkent_equivalent.course_type == MUST_COURSE:  # bilkent course name
                    courses_table.cell(2 + i, 4).text = bilkent_equivalent.course_code + " " + \
                                                        bilkent_equivalent.course_name
                elif bilkent_equivalent.course_type == ELECTIVE_COURSE:
                    courses_table.cell(2 + i, 4).text = bilkent_equivalent.elective_group_name
                courses_table.cell(2 + i, 5).text = str(bilkent_equivalent.course_credit)  # bilkent course credit
                if bilkent_equivalent.course_type == ELECTIVE_COURSE:  # bilkent elective course code
                    courses_table.cell(2 + i, 6).text = bilkent_equivalent.course_code

        for merged_course, course_list in user_merged_course_dict.items():
            for i, user_course in enumerate(course_list):
                j = i + len(unmerged_courses)
                course = user_course.course
                bilkent_equivalent = course.bilkent_equivalent
                courses_table.cell(2 + j, 1).text = course.code  # course code
                courses_table.cell(2 + j, 2).text = course.course_name  # course name
                courses_table.cell(2 + j, 3).text = str(course.course_credit)  # course credit
                if bilkent_equivalent is not None:
                    if bilkent_equivalent.course_type == MUST_COURSE:  # bilkent course name
                        courses_table.cell(2 + j, 4).text = bilkent_equivalent.course_code + " " + \
                                                            bilkent_equivalent.course_name
                    elif bilkent_equivalent.course_type == ELECTIVE_COURSE:
                        courses_table.cell(2 + j, 4).text = bilkent_equivalent.elective_group_name
                    courses_table.cell(2 + j, 5).text = str(bilkent_equivalent.course_credit)  # bilkent course credit
                    if bilkent_equivalent.course_type == ELECTIVE_COURSE:  # bilkent elective course code
                        courses_table.cell(2 + j, 6).text = bilkent_equivalent.course_code
                if i is not len(course_list)-1:
                    courses_table.cell(2+j, 4).merge(courses_table.cell(3+j, 4))
                    courses_table.cell(2 + j, 5).merge(courses_table.cell(3 + j, 5))
                    courses_table.cell(2 + j, 6).merge(courses_table.cell(3 + j, 6))


        coordinator_table.cell(1, 1).text = student.coordinator.user.name  # coordinator name

        turkey_timezone = pytz.timezone("Turkey")
        date = datetime.now(turkey_timezone)
        document_name = STATIC_DOCUMENTS_FOLDER + 'pre_approval_form' + date.strftime("%d-%m-%Y-%H-%M-%S") + ".docx"
        document.save(document_name)

        return document_name, date, PREAPPROVAL_FORM

class CreateLearningAgreementView(CreateDocumentView):
    def fill_necessary_information(self, student):
        document = docx.api.Document('courses/static/Learning_Agreement.docx')

        student_info_table = document.tables[0]
        host_uni_tableB = document.tables[1]

        student_info_table.cell(1, 1).text = student.user.name.split()[:-1]  # name
       # student_info_table.cell(1, 3).text = str(student.user.bilkent_id)  # bilkent id number
        student_info_table.cell(1, 2).text = student.user.name.split()[-1]  # surname
        student_info_table.cell(3, 4).text = student.user.department  # department
        student_info_table.cell(3, 1).text = student.university.university_name  # university


        document_name = STATIC_DOCUMENTS_FOLDER + 'learning_agreement_form.docx'
        document.save(document_name)

        return document_name, "", PREAPPROVAL_FORM



import aspose.words as aw
from datetime import date


class CompareDocument(View):

    def get(self, request):
        # load first document
        doc = aw.Document("static/documents/documents/pre_approval_form17-12-2022-13-32-48.docx")

        # load second document
        doc2 = aw.Document("static/documents/documents/pre_approval_form17-12-2022-13-58-08.docx")

        # compare documents
        doc.compare(doc2, "user", date.today())

        # save the document to get the revisions
        if (doc.revisions.count > 0):
            print("There are some differences between documents")
            doc.save("static/documents/documents/compared.docx")
        else:
            print("Documents are equal")

        messages.success("Documents are compared successfully!!")
        return redirect('/documents')


# compare two documents to see the differences
class CompareDocuments(LoginRequiredMixin, View):
    def get(self, request, student_id, doc1_id, doc2_id):
        user = request.user
        erasmus_user = ErasmusUser.objects.filter(user=user).first()
        coordinator = Coordinator.objects.filter(user=erasmus_user).first()
        if coordinator is None:
            student_request = Student.objects.filter(user=erasmus_user).first()

        student = get_object_or_404(Student, pk=student_id)
        if student.user.name != student.user.name:
            messages.error(request, 'you are not allowed to compare documents')
            return redirect('/documents')

        document1 = get_object_or_404(Document, pk=doc1_id)
        document2 = get_object_or_404(Document, pk=doc2_id)

        print(document1.user.user.name)
        print(document2.user.user.name)
        if document1.user.user.name != document2.user.user.name:
            messages.error(request, "You can only compare documents which belongs to same user")
            return redirect('/documents')

        if document1.document_type != document2.document_type:
            messages.error(request, "You can only compare documents having same document type")
            return redirect('/documents')

        doc1 = aw.Document(document1.document.path)

        doc2 = aw.Document(document2.document.path)

        doc2.compare(doc1, "user", date.today())

        if doc2.revisions.count > 0:
            document_name = STATIC_DOCUMENTS_FOLDER + 'compared.docx'
            doc2.save(document_name)
            new_document_name = 'Compared ' + document1.document_type + ' ' + document1.user.user.name + str(document1.date.strftime('%m/%d/%Y %H:%M'))\
                                + ' and ' + str(document2.date.strftime('%m/%d/%Y %H:%M'))

            with open(document_name, 'rb') as f:
                new_document = Document(document_name=new_document_name, user=student, document_type=document1.document_type)
                new_document.document = File(f, name=os.path.basename(f.name))
                new_document.save()
            messages.success(request, "Documents are compared")
            return redirect('/documents')

        else:
            print("Documents are equal")
            return redirect('/documents')


